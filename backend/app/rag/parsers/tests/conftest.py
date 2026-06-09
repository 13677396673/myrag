"""文档解析器测试夹具

提供：
- 文本夹具文件的路径
- 自动生成的 sample.pdf 和 sample.docx 夹具
- 已注册默认解析器的 ParserRouter 实例
"""

import os
from pathlib import Path

import pytest

from app.rag.parsers import ParserRouter, register_default_parsers

FIXTURES_DIR = Path(__file__).parent / "fixtures"


# ── 文本夹具路径 ──────────────────────────────────────────────


@pytest.fixture
def sample_txt_path() -> str:
    """返回 ``sample.txt`` 的绝对路径"""
    return str(FIXTURES_DIR / "sample.txt")


@pytest.fixture
def sample_md_path() -> str:
    """返回 ``sample.md`` 的绝对路径"""
    return str(FIXTURES_DIR / "sample.md")


@pytest.fixture
def sample_pdf_path(tmp_path: Path) -> str:
    """使用 PyMuPDF 动态生成一个 3 页的 PDF 文件用于测试"""
    import fitz

    pdf_path = tmp_path / "sample.pdf"
    doc = fitz.open()

    # 第 1 页
    page1 = doc.new_page()
    page1.insert_text((72, 72), "Page 1 Content Here", fontsize=12)

    # 第 2 页
    page2 = doc.new_page()
    page2.insert_text((72, 72), "Page 2 Content Here", fontsize=12)

    # 第 3 页
    page3 = doc.new_page()
    page3.insert_text((72, 72), "Page 3 Content Here", fontsize=12)

    doc.save(str(pdf_path))
    doc.close()
    return str(pdf_path)


@pytest.fixture
def sample_docx_path(tmp_path: Path) -> str:
    """使用 python-docx 动态生成一个 .docx 文件用于测试"""
    from docx import Document

    docx_path = tmp_path / "sample.docx"
    doc = Document()

    doc.add_heading("Test Document", level=1)
    doc.add_paragraph("This is a test paragraph for DOCX parsing.")
    doc.add_paragraph("这是第二个段落，包含中文内容。")
    doc.add_paragraph("Third paragraph with mixed English and 中文内容.")

    doc.save(str(docx_path))
    return str(docx_path)


@pytest.fixture
def empty_txt_path(tmp_path: Path) -> str:
    """返回一个空文本文件的路径"""
    empty_path = tmp_path / "empty.txt"
    empty_path.write_text("", encoding="utf-8")
    return str(empty_path)


# ── ParserRouter 夹具 ─────────────────────────────────────────


@pytest.fixture
def parser_router() -> ParserRouter:
    """返回一个已注册默认解析器的 ParserRouter 实例"""
    router = ParserRouter()
    register_default_parsers(router)
    return router

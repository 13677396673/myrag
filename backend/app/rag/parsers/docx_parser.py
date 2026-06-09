"""DOCX 解析器 — DocxParser

基于 python-docx 提取 ``.docx`` 文件中的段落文本。

用法::

    parser = DocxParser()
    doc = parser.parse("report.docx")
"""

from typing import List

from docx import Document as DocxDocument

from app.rag.interfaces.parser import DocumentParser, ParsedDocument


class DocxParser(DocumentParser):
    """Word 文档解析器（依赖 python-docx）"""

    def parse(self, file_path: str) -> ParsedDocument:
        """解析 .docx 文件

        提取所有段落文本，用 ``\\n\\n`` 连接。

        参数:
            file_path: .docx 文件的完整路径

        返回:
            ParsedDocument，包含完整文本

        异常:
            FileNotFoundError: 文件不存在
        """
        doc = DocxDocument(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        content = "\n\n".join(paragraphs)

        return ParsedDocument(
            content=content,
            metadata={"source": file_path, "paragraphs": len(paragraphs)},
        )

    @classmethod
    def supported_extensions(cls) -> List[str]:
        return [".docx"]
